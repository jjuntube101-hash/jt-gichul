# -*- coding: utf-8 -*-
"""해설검수 샘플 50문항 MD -> Word 변환"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    elm = shading.makeelement(qn("w:shd"), {qn("w:fill"): color, qn("w:val"): "clear"})
    shading.append(elm)


# MD 파일 파싱
with open("해설검수_샘플50문항.md", "r", encoding="utf-8") as f:
    content = f.read()

# 테이블 행 추출
table_rows = []
for line in content.split("\n"):
    m = re.match(
        r"\|\s*(\d+)\s*\|\s*(\d+)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(\d+)\s*\|\s*(\d/5)\s*\|",
        line,
    )
    if m:
        table_rows.append(m.groups())

print(f"파싱된 문항: {len(table_rows)}건")

# Word 문서 생성
doc = Document()

style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

# 제목
title = doc.add_heading("해설 품질 검수용 샘플 50문항", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("랜덤 추출 (과목별 균등 배분). 해설의 정확성과 완성도를 확인해주세요.")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("URL 열의 주소를 브라우저에 입력하면 해설 페이지로 이동합니다.")
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(150, 150, 150)

doc.add_paragraph()

# 메인 테이블
headers = ["#", "no", "과목", "중분류", "시험", "연도", "난이도", "URL", "검수"]
table = doc.add_table(rows=1 + len(table_rows), cols=len(headers))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 헤더
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, "4F46E5")

# 데이터
for row_idx, (num, no, subj, topic, exam, year, diff) in enumerate(table_rows):
    row = table.rows[row_idx + 1]
    url = f"gichul.jttax.co.kr/question/{no}"
    values = [num, no, subj.strip(), topic.strip(), exam.strip(), year, diff, url, ""]

    for col_idx, val in enumerate(values):
        cell = row.cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
            if col_idx in [0, 1, 5, 6, 8]:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if row_idx % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, "F1F5F9")

# 컬럼 너비
widths = [Cm(0.8), Cm(1), Cm(2.2), Cm(3.2), Cm(1.8), Cm(1), Cm(1), Cm(4.2), Cm(1.5)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

doc.add_paragraph()

# 검수 체크리스트
doc.add_heading("검수 체크리스트", level=2)
p = doc.add_paragraph("각 문항별로 아래 항목을 확인해주세요:")

checklist = [
    "선지별 O/X 판정 정확성",
    "근거 조문 정확성 (법령/조/항/호)",
    "함정 유형 분류 적절성",
    "출제의도 설명 적절성",
    "학습 가이드 유용성",
    "오탈자 여부",
]
for item in checklist:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(10)

doc.add_paragraph()

# 검수 메모란
doc.add_heading("검수 메모", level=2)
p = doc.add_paragraph("오류나 개선사항을 아래에 기록해주세요:")
p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

memo = doc.add_table(rows=11, cols=2)
memo.style = "Table Grid"
memo.rows[0].cells[0].text = "문항 no"
memo.rows[0].cells[1].text = "메모 (오류/개선사항)"
for cell in memo.rows[0].cells:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)
    set_cell_shading(cell, "E2E8F0")

for r in memo.rows:
    r.cells[0].width = Cm(2.5)
    r.cells[1].width = Cm(14)

out = "해설검수_샘플50문항.docx"
doc.save(out)
print(f"완료: {out}")
